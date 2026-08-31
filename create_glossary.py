from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('AIIA-CTMS Technical Glossary', 0)
doc.add_paragraph('A comprehensive reference of all technical terms, concepts, and acronyms used in the SIH 2026 Clinical Trials Management System project.').paragraph_format.space_after = Pt(12)

# Regulatory & Compliance
doc.add_heading('1. Regulatory & Compliance Terms', 1)
glossary = [
    ('CTRI', 'Clinical Trials Registry - India. The official Indian registry for clinical trials; required for trial registration in India.'),
    ('NDCT Rules 2019', 'National Database of Clinical Trials Rules, 2019. Defines regulatory requirements and timelines for clinical trial reporting and AE/SAE escalation in India.'),
    ('GCP', 'Good Clinical Practice. International standard for ethical and scientific quality requirements for designing, conducting, recording, and reporting trials.'),
    ('GCP-ASU', 'Good Clinical Practice guidelines specific to Ayurveda, Siddha, and Unani systems of medicine as per ICMR standards.'),
    ('IEC', 'Institutional Ethics Committee. The body responsible for reviewing and approving clinical trial protocols, amendments, and adverse event reports.'),
    ('ICMR', 'Indian Council of Medical Research. A premier research body that sets guidelines for clinical research in India.'),
    ('ICH Guidelines', 'International Council for Harmonisation guidelines. Global standards for clinical trial conduct and reporting.'),
    ('SAE', 'Serious Adverse Event. An unintended harmful reaction to a study drug that is serious (hospitalization, disability, death, etc.).'),
    ('AE', 'Adverse Event. Any undesirable experience by a subject during a clinical trial, regardless of whether it is related to the study drug.'),
    ('Causality Assessment', 'Determination of the relationship between a drug and an adverse event (Certain, Probable, Possible, Unlikely, or Unrelated).'),
    ('Data Minimization', 'Security principle of collecting and storing only the minimum personal data necessary for a specific purpose.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

# Clinical Research
doc.add_heading('2. Clinical Research & Trial Management', 1)
glossary = [
    ('Study / Trial', 'A research investigation on human subjects following a predetermined protocol to evaluate safety and/or efficacy of an intervention.'),
    ('Principal Investigator (PI)', 'The lead researcher responsible for conducting a trial at a site, enrolling subjects, and collecting data.'),
    ('Subject / Participant', 'An individual enrolled in a clinical trial. De-identified using system-generated unique IDs to protect privacy.'),
    ('De-identification', 'Removal or encryption of personally identifiable information (PII) to protect subject privacy; data is identified only by a unique system ID.'),
    ('Enrollment Status', 'Current status of a subject in the trial: Screened, Randomized, Completed, or Withdrawn.'),
    ('Screening', 'Initial assessment of potential subjects to determine eligibility for trial participation.'),
    ('Randomization', 'Process of assigning subjects to treatment groups in a way determined by chance, not by any deliberate decision.'),
    ('Consent / Informed Consent', 'Documented agreement by a subject to participate in a trial after being fully informed of its nature, risks, and benefits.'),
    ('Protocol', 'Detailed plan describing all aspects of a clinical trial, including objectives, design, methodology, statistical analysis, and ethical considerations.'),
    ('Protocol Deviation', 'Unintended or unapproved deviation from the study protocol; must be documented and reviewed by IEC.'),
    ('Milestone', 'Key event or deadline in the trial timeline (e.g., protocol approval, enrollment start, interim analysis, trial closure).'),
    ('Therapeutic Area', 'Medical specialty or disease category for the trial (e.g., Cardiology, Rheumatology, Gastrointestinal).'),
    ('Dosha Profile', 'Ayurveda-specific categorization of a subjects constitution (Vata, Pitta, Kapha) used for baseline characterization.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

# CDISC Standards
doc.add_heading('3. CDISC Data Standards', 1)
glossary = [
    ('CDISC', 'Clinical Data Interchange Standards Consortium. International organization setting standards for clinical research data.'),
    ('CDASH', 'Clinical Data Acquisition Standards Harmonization. Standard for electronic case report form (eCRF) field design and data capture.'),
    ('SDTM', 'Study Data Tabulation Model. Standard for organizing and structuring study data for regulatory submission.'),
    ('SDTM Domains', 'Categories of data in the SDTM structure: DM (Demographics), AE (Adverse Events), VS (Vital Signs), CM (Concomitant Medications), etc.'),
    ('DM Domain', 'Demographics domain in SDTM; contains subject baseline characteristics like age, sex, race, enrollment status.'),
    ('AE Domain', 'Adverse Events domain in SDTM; contains all recorded adverse events and their attributes.'),
    ('VS Domain', 'Vital Signs domain in SDTM; contains physiological measurements like blood pressure, heart rate, temperature.'),
    ('CM Domain', 'Concomitant Medications domain in SDTM; contains non-study medications taken by subjects during the trial.'),
    ('USUBJID', 'Unique Subject Identifier; the unique de-identified ID assigned to each subject in the study.'),
    ('Define-XML', 'Standard for documenting metadata about SDTM datasets, including variable definitions, value lists, and computations for regulatory submission.'),
    ('eCRF', 'Electronic Case Report Form. Digital form for capturing study data at the point of care, based on CDASH standards.'),
    ('Data Query', 'Request from data manager to site or investigator to clarify, verify, or correct submitted data.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

# Technical Architecture
doc.add_heading('4. Technical Architecture & Infrastructure', 1)
glossary = [
    ('REST API', 'Representational State Transfer API. Web service that uses HTTP methods (GET, POST, PUT, DELETE) to perform CRUD operations.'),
    ('OAuth2', 'Industry-standard authorization protocol allowing users to authenticate and grant permissions without sharing passwords.'),
    ('OpenID Connect', 'Authentication layer built on top of OAuth2, used for user identity verification.'),
    ('JWT', 'JSON Web Token. Compact token format for securely transmitting claims between parties; used for session management.'),
    ('MFA', 'Multi-Factor Authentication. Security mechanism requiring two or more verification methods (e.g., password + TOTP code).'),
    ('TOTP', 'Time-based One-Time Password. A type of two-factor authentication that generates temporary codes every 30 seconds.'),
    ('RBAC', 'Role-Based Access Control. Security model restricting system access based on user roles (PI, IEC, DataManager, etc.).'),
    ('TLS 1.2+', 'Transport Layer Security. Encryption protocol for secure communication over HTTPS; version 1.2 or higher for modern security standards.'),
    ('AES-256', 'Advanced Encryption Standard with 256-bit key length. Encryption algorithm for protecting sensitive data at rest.'),
    ('Parameterized Queries', 'SQL queries with placeholders for values, preventing SQL injection attacks by separating code from data.'),
    ('PostgreSQL', 'Open-source relational database system; primary data store for the AIIA-CTMS backend.'),
    ('S3-Compatible Storage', 'Object storage service compatible with Amazon S3 API for storing documents, PDFs, and consent forms.'),
    ('Docker', 'Containerization technology for packaging applications and dependencies in isolated containers.'),
    ('Cloud Hosting', 'Deployment of application on cloud platforms (GCP, Azure) with automatic scaling and managed services.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

# FHIR & Interoperability
doc.add_heading('5. FHIR & Interoperability', 1)
glossary = [
    ('FHIR', 'Fast Healthcare Interoperability Resources (HL7 FHIR R4). Modern standard for exchanging healthcare data across systems.'),
    ('FHIR R4', 'Release 4 of FHIR; the current stable version with comprehensive resource definitions for healthcare entities.'),
    ('HAPI FHIR', 'Open-source Java framework for building and consuming FHIR APIs; used as the FHIR gateway in AIIA-CTMS.'),
    ('Patient Resource', 'FHIR resource representing demographic and administrative information about a subject/patient.'),
    ('AdverseEvent Resource', 'FHIR resource representing an adverse event occurrence with details on causality, severity, and outcome.'),
    ('ResearchStudy Resource', 'FHIR resource representing a clinical research study with metadata on objectives, phases, and participants.'),
    ('Location Resource', 'FHIR resource representing a physical site or facility (e.g., research center location).'),
    ('Organization Resource', 'FHIR resource representing an organization (e.g., Ministry of Ayush, AIIA, hospital).'),
    ('ABDM', 'Ayushman Bharat Digital Mission. Indias health data exchange framework; AIIA-CTMS has stub integration for future linking.'),
    ('Health ID', 'Unique identifier in the ABDM ecosystem linking a citizen to their electronic health records.'),
    ('Interoperability', 'Ability of different systems to exchange and use data seamlessly; FHIR enables this across healthcare IT systems.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

# Database & Data
doc.add_heading('6. Database & Data Management', 1)
glossary = [
    ('Entity-Relationship Model', 'Database design approach using entities (tables) and their relationships (foreign keys) to structure data.'),
    ('Primary Key (PK)', 'Unique identifier for each record in a table; ensures uniqueness and enables fast lookups.'),
    ('Foreign Key (FK)', 'Reference to a primary key in another table, establishing relationships between entities.'),
    ('Schema', 'Logical blueprint of a database defining tables, columns, data types, constraints, and relationships.'),
    ('Audit Log', 'Tamper-evident record of all changes to data, including who made the change, when, what changed, and before/after values.'),
    ('Hash Chain', 'Cryptographic technique linking audit records by including a hash of the previous record in the current one, detecting tampering.'),
    ('Record Hash', 'Cryptographic hash computed from audit record data plus the previous record hash; used for tamper-evidence.'),
    ('Immutable Audit Trail', 'Audit log that cannot be modified or deleted after creation, ensuring integrity and regulatory compliance.'),
    ('Data Completeness', 'Percentage of required fields populated for a subject or dataset; tracked per SDTM domain.'),
    ('Data Quality', 'Assessment of whether data is accurate, complete, consistent, and compliant with CDISC standards.'),
    ('Compliance Score', 'Computed metric reflecting adherence to regulatory checkpoints and GCP requirements for a study.'),
    ('Indexing', 'Database optimization technique creating fast lookup structures on frequently queried columns (study_id, subject_id, event_date, etc.).'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

# Security & Audit
doc.add_heading('7. Security, Compliance & Audit Concepts', 1)
glossary = [
    ('Tamper Detection', 'Capability to identify unauthorized or unintended modifications to data using cryptographic hashing and audit trails.'),
    ('Tamper-Evident', 'System design that makes any unauthorized changes immediately detectable; the audit chain breaks if data is modified.'),
    ('PII', 'Personally Identifiable Information. Sensitive data about subjects (name, address, contact) that must be protected and minimized.'),
    ('Subject PII Table', 'Separate database table containing encrypted personally identifiable information, accessed only by authorized roles.'),
    ('Escalation', 'Process of automatically increasing alert level or notification recipients when a critical event (e.g., SAE with missed deadline) occurs.'),
    ('Escalation Level', 'Tier in escalation hierarchy: Level 0 (PI), Level 1 (IEC), Level 2 (Sponsor), Level 3 (NPvCC/Ministry).'),
    ('Deadline Watcher', 'Background job running periodically to check open AE/SAE events against reporting deadlines and trigger escalations.'),
    ('Rules Engine', 'Configurable system that applies regulatory rules (e.g., SAE initial report within 24 hours) to compute deadlines and actions.'),
    ('Permission Matrix', 'Table defining which roles have which permissions (Create, Read, Update, Delete) for each resource type.'),
    ('Session Management', 'Mechanism for tracking authenticated users; typically using JWT tokens or server-side session stores.'),
    ('Data Residency', 'Requirement that data be stored within a specific geographic region (e.g., India) for regulatory compliance.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

# System Roles
doc.add_heading('8. System Roles & User Types', 1)
glossary = [
    ('PI (Principal Investigator)', 'Site-level role responsible for trial conduct, subject enrollment, data entry, and AE/SAE reporting.'),
    ('IEC Member', 'Ethics committee role reviewing protocols, amendments, and adverse events for regulatory approval.'),
    ('Data Manager', 'Responsible for data quality, CDASH/SDTM compliance, query resolution, and SDTM dataset export.'),
    ('Pharmacovigilance Officer / PV Officer', 'Monitors AE/SAE signals nationally (NPvCC), ensures reporting timeline compliance, tracks escalations.'),
    ('Sponsor / Trial Admin', 'Portfolio-level role overseeing multiple trials, budgets, timelines, and compliance status.'),
    ('Executive Viewer', 'Ministry/leadership role viewing aggregate national KPIs, dashboards, and strategic metrics.'),
    ('Auditor / Compliance Officer', 'Reviews audit trail, data integrity, chain-of-custody, and prepares for regulatory inspections.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

# Acronyms & Abbreviations
doc.add_heading('9. Common Acronyms & Abbreviations', 1)
glossary = [
    ('API', 'Application Programming Interface'),
    ('SMTP', 'Simple Mail Transfer Protocol'),
    ('SMS', 'Short Message Service'),
    ('JSON', 'JavaScript Object Notation'),
    ('XML', 'Extensible Markup Language'),
    ('CSV', 'Comma-Separated Values'),
    ('HTTPS', 'HTTP Secure (encrypted)'),
    ('OAuth', 'Open Authorization'),
    ('JWT', 'JSON Web Token'),
    ('RBAC', 'Role-Based Access Control'),
    ('PII', 'Personally Identifiable Information'),
    ('CRUD', 'Create, Read, Update, Delete'),
    ('MVP', 'Minimum Viable Product'),
    ('KPI', 'Key Performance Indicator'),
    ('ETL', 'Extract, Transform, Load'),
    ('FHIR', 'Fast Healthcare Interoperability Resources'),
    ('HL7', 'Health Level 7'),
    ('SNOMED', 'Systematized Nomenclature of Medicine'),
    ('UUID', 'Universally Unique Identifier'),
    ('ISO8601', 'International standard for date/time format'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

# Key Concepts
doc.add_heading('10. Key Project Concepts', 1)
glossary = [
    ('Dashboard', 'Role-specific UI view providing personalized information, metrics, and actions for each user role.'),
    ('Multi-Centre Trial', 'A clinical trial conducted at multiple research sites under one protocol, overseen by one sponsor.'),
    ('Continuing Review', 'Periodic IEC review (typically annual) of an ongoing trial to reassess risk/benefit and ensure compliance.'),
    ('Reporting Deadline', 'Regulatory time limit for reporting an AE/SAE to authorities (e.g., 24 hours for SAE initial report).'),
    ('Study Status', 'Current phase of a trial: Draft, Pending IEC Approval, Active, Closed, or Suspended.'),
    ('Subject Enrollment Funnel', 'Visual representation of subject progression: Screened -> Randomized -> Completed -> Withdrawn.'),
    ('Synthetic Data', 'Artificially generated data used for testing/demo purposes (no real subject information).'),
    ('Rule Configuration', 'JSON-based settings defining regulatory logic (timelines, thresholds) that can be updated without code changes.'),
    ('Background Job', 'Automated task running periodically (cron) to process data, check conditions, or trigger actions (e.g., deadline watcher).'),
    ('Data Export', 'Download of study data in standardized formats (SDTM, Define-XML, CSV, PDF) for regulatory submission or archival.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(term).bold = True
    p.add_run(': ' + definition)

doc.add_paragraph('')
doc.add_paragraph('This glossary is a reference document for understanding technical concepts, standards, and terminology used throughout the AIIA-CTMS project. For additional context, refer to the PRD, TRD, Architecture Guide, and API documentation.', style='Normal')

doc.save('AIIA_CTMS_Technical_Glossary.docx')
print('Created: AIIA_CTMS_Technical_Glossary.docx')
